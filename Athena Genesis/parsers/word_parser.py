# -*- coding: utf-8 -*-
"""
Word文件解析器
"""

import os
import docx
from typing import List, Dict, Any
from .base_parser import BaseParser
from config.settings import SETTINGS


class WordParser(BaseParser):
    """Word文件解析器"""

    def __init__(self):
        super().__init__()
        self.supported_extensions = SETTINGS.SUPPORTED_EXTENSIONS['word']
        self.parser_name = "WordParser"
        self.max_paragraphs_to_show = 100

    def can_parse(self, file_path: str) -> bool:
        """检查是否能解析Word文件"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext == '.docx'

    def parse(self, file_path: str) -> str:
        """解析Word文档"""
        try:
            doc = docx.Document(file_path)

            content_lines = [f"[Word文档: {os.path.basename(file_path)}]"]
            content_lines.append(f"段落数: {len(doc.paragraphs)}")

            # 解析段落
            paragraph_count = 0
            for i, para in enumerate(doc.paragraphs):
                if paragraph_count >= self.max_paragraphs_to_show:
                    remaining = len(doc.paragraphs) - i
                    content_lines.append(f"\n...省略 {remaining} 个段落")
                    break

                text = para.text.strip()
                if text:
                    content_lines.append(text)
                    paragraph_count += 1

            # 解析表格
            if doc.tables:
                content_lines.append(f"\n[发现 {len(doc.tables)} 个表格]")
                for table_num, table in enumerate(doc.tables[:3]):  # 最多3个表格
                    content_lines.append(f"\n表格 {table_num + 1}:")
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        if row_text.strip():
                            content_lines.append(row_text)

            return "\n".join(content_lines)

        except Exception as e:
            return f"[Word解析错误]: {str(e)}"

    def get_document_structure(self, file_path: str) -> Dict[str, Any]:
        """获取文档结构"""
        try:
            doc = docx.Document(file_path)

            structure = {
                "filename": os.path.basename(file_path),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "heading_count": 0,
                "headings": [],
                "styles": set(),
                "sections": []
            }

            # 分析标题
            for para in doc.paragraphs:
                style_name = para.style.name
                structure["styles"].add(style_name)

                if style_name.startswith('Heading'):
                    structure["heading_count"] += 1
                    structure["headings"].append({
                        "text": para.text[:50],
                        "style": style_name,
                        "level": int(style_name.replace('Heading ', '')) if style_name.replace('Heading ',
                                                                                               '').isdigit() else 0
                    })

            structure["styles"] = list(structure["styles"])

            return structure

        except Exception as e:
            return {"error": str(e)}

    def extract_images(self, file_path: str, output_dir: str) -> List[str]:
        """提取Word中的图片"""
        try:
            import zipfile
            import shutil

            os.makedirs(output_dir, exist_ok=True)
            image_paths = []

            # Word文档实际上是ZIP文件
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                # 提取所有图片文件
                image_files = [f for f in zip_ref.namelist()
                               if f.startswith('word/media/')]

                for image_file in image_files:
                    # 提取文件
                    image_name = os.path.basename(image_file)
                    output_path = os.path.join(output_dir, image_name)

                    with zip_ref.open(image_file) as source, \
                            open(output_path, 'wb') as target:
                        shutil.copyfileobj(source, target)

                    image_paths.append(output_path)

            return image_paths

        except Exception as e:
            print(f"提取图片失败: {e}")
            return []